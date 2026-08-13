from docx import Document
from docx.shared import RGBColor
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy import select
from sqlalchemy.orm import Session
from sqlalchemy.pool import StaticPool

from app.ai.llm import LlmClient
from app.db.base import Base
from app.db.session import get_db
from app.db.models import Chapter, ChapterDraft, ChapterRelation, PaperAbstract, Project
from app.export.docx import build_docx
from app.export.docx_style import DocxStyle
from app.export.markdown import build_markdown
from app.main import create_app
from app.services.chapters import list_chapters_in_hierarchy_order
from app.services.generation import GenerationService


class NestedOutlineClient(LlmClient):
    def complete_json(self, prompt: str) -> dict:
        return {
            "chapters": [
                {
                    "title": "Second root",
                    "level": 1,
                    "order": 2,
                    "children": [],
                },
                {
                    "title": "First root",
                    "level": 1,
                    "order": 1,
                    "children": [
                        {
                            "title": "Second child",
                            "level": 2,
                            "order": 2,
                            "children": [],
                        },
                        {
                            "title": "First child",
                            "level": 2,
                            "order": 1,
                            "children": [
                                {
                                    "title": "Grandchild",
                                    "level": 3,
                                    "order": 1,
                                    "children": [],
                                }
                            ],
                        },
                    ],
                },
            ]
        }

    def complete_markdown(self, prompt: str) -> str:
        raise AssertionError("Outline should not request Markdown")


def test_nested_outline_generation_and_listing_use_hierarchy_order(db_session):
    project = Project(type="thesis", title="Nested", language="zh")
    db_session.add(project)
    db_session.commit()

    generated = GenerationService.generate_outline(
        db_session, project.id, client=NestedOutlineClient()
    )
    listed = list_chapters_in_hierarchy_order(db_session, project.id)

    expected = [
        "First root",
        "First child",
        "Grandchild",
        "Second child",
        "Second root",
    ]
    assert [chapter.title for chapter in generated] == expected
    assert [chapter.title for chapter in listed] == expected


def test_markdown_export_uses_hierarchy_order_without_duplicate_heading(db_session):
    project = Project(type="thesis", title="Nested export", language="zh")
    root = Chapter(project=project, title="Root", level=1, order=1)
    child = Chapter(project=project, parent=root, title="Child", level=2, order=1)
    draft = ChapterDraft(
        chapter=child,
        version=1,
        content="# Child\n\nBody",
        generation_mode="generate",
    )
    db_session.add_all([project, root, child, draft])
    db_session.commit()

    chapters = list_chapters_in_hierarchy_order(db_session, project.id)
    markdown = build_markdown(project, chapters, {child.id: draft})

    assert markdown.index("# Root") < markdown.index("## Child")
    assert markdown.count("## Child") == 1
    assert "# Child" not in markdown.replace("## Child", "")


def test_markdown_export_skips_child_headings_already_in_parent_draft(db_session):
    project = Project(type="thesis", title="Hive 电商数仓", language="zh")
    parent = Chapter(project=project, title="绪论", level=1, order=1)
    background = Chapter(
        project=project, parent=parent, title="项目背景与意义", level=2, order=1
    )
    status = Chapter(
        project=project, parent=parent, title="国内外研究现状", level=2, order=2
    )
    structure = Chapter(
        project=project, parent=parent, title="主要研究内容与论文结构", level=2, order=3
    )
    next_chapter = Chapter(project=project, title="相关技术基础", level=1, order=2)
    parent_draft = ChapterDraft(
        chapter=parent,
        version=1,
        content=(
            "# 第1章 绪论\n\n"
            "引言。\n\n"
            "## 1.1 项目背景与意义\n\n"
            "背景段落。\n\n"
            "## 1.2 国内外研究现状\n\n"
            "现状段落。\n\n"
            "## 1.3 主要研究内容与论文结构\n\n"
            "结构段落。\n\n"
            "## 1.4 本章小结\n\n"
            "本章介绍了电商数据仓库背景。\n"
        ),
        generation_mode="generate",
    )
    next_draft = ChapterDraft(
        chapter=next_chapter,
        version=1,
        content="# 第2章 相关技术基础\n\nHive 与 Hadoop。\n",
        generation_mode="generate",
    )
    db_session.add_all(
        [
            project,
            parent,
            background,
            status,
            structure,
            next_chapter,
            parent_draft,
            next_draft,
        ]
    )
    db_session.commit()

    chapters = list_chapters_in_hierarchy_order(db_session, project.id)
    markdown = build_markdown(
        project,
        chapters,
        {parent.id: parent_draft, next_chapter.id: next_draft},
    )

    assert markdown.count("项目背景与意义") == 1
    assert markdown.count("国内外研究现状") == 1
    assert markdown.count("主要研究内容与论文结构") == 1
    assert "## 项目背景与意义" not in markdown
    assert "## 国内外研究现状" not in markdown
    assert "## 主要研究内容与论文结构" not in markdown
    assert "# 相关技术基础" in markdown
    assert "Hive 与 Hadoop。" in markdown
    assert markdown.index("本章小结") < markdown.index("相关技术基础")


def test_docx_export_skips_child_headings_already_in_parent_draft(db_session, tmp_path):
    project = Project(type="thesis", title="Hive 电商数仓", language="zh")
    parent = Chapter(project=project, title="绪论", level=1, order=1)
    child = Chapter(project=project, parent=parent, title="项目背景与意义", level=2, order=1)
    next_chapter = Chapter(project=project, title="相关技术基础", level=1, order=2)
    parent_draft = ChapterDraft(
        chapter=parent,
        version=1,
        content="# 第1章 绪论\n\n## 1.1 项目背景与意义\n\n背景段落。\n\n## 1.4 本章小结\n\n小结。\n",
        generation_mode="generate",
    )
    next_draft = ChapterDraft(
        chapter=next_chapter,
        version=1,
        content="# 第2章 相关技术基础\n\nHive。\n",
        generation_mode="generate",
    )
    db_session.add_all([project, parent, child, next_chapter, parent_draft, next_draft])
    db_session.commit()
    output = tmp_path / "export.docx"
    chapters = list_chapters_in_hierarchy_order(db_session, project.id)

    build_docx(output, project, chapters, {parent.id: parent_draft, next_chapter.id: next_draft})

    headings = [
        paragraph.text
        for paragraph in Document(output).paragraphs
        if paragraph.style.name.startswith("Heading")
    ]
    assert "项目背景与意义" not in headings
    assert "绪论" in headings
    assert "相关技术基础" in headings


def test_docx_export_uses_academic_styles_instead_of_markdown_markers(tmp_path):
    project = Project(type="thesis", title="基于Hive的电商数据仓库设计与实现", language="zh")
    chapter = Chapter(project=project, title="绪论", level=1, order=1)
    draft = ChapterDraft(
        chapter=chapter,
        version=1,
        content=(
            "# 第1章 绪论\n\n"
            "## 1.1 项目背景与意义\n\n"
            "本章介绍**研究背景**，详见[资料](https://example.com)。\n\n"
            "- **要点**一项\n"
        ),
        generation_mode="generate",
    )
    output = tmp_path / "styled.docx"

    build_docx(
        output,
        project,
        [chapter],
        {chapter.id: draft},
        DocxStyle(body_size_pt=12, heading_east_asia="黑体", body_east_asia="宋体"),
    )
    document = Document(output)
    texts = [paragraph.text for paragraph in document.paragraphs]
    joined = "\n".join(texts)
    assert "##" not in joined
    assert "**" not in joined
    assert "](" not in joined
    assert "1.1 项目背景与意义" in texts
    assert "本章介绍研究背景，详见资料。" in texts
    assert "• 要点一项" in texts
    heading_texts = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.style.name.startswith("Heading")
    ]
    assert "绪论" in heading_texts
    assert "1.1 项目背景与意义" in heading_texts
    body = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text == "本章介绍研究背景，详见资料。"
    )
    assert body.runs[0].font.size.pt == 12
    assert body.runs[0]._element.rPr.rFonts.get(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia"
    ) == "宋体"
    assert body.paragraph_format.line_spacing == 1.5
    assert body.paragraph_format.first_line_indent.pt == 24
    bold_runs = [run for run in body.runs if run.text == "研究背景"]
    assert bold_runs and bold_runs[0].bold
    h1 = next(paragraph for paragraph in document.paragraphs if paragraph.text == "绪论")
    assert h1.runs[0].font.size.pt == 15
    assert h1.runs[0]._element.rPr.rFonts.get(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia"
    ) == "黑体"
    h1_style = document.styles["Heading 1"]
    assert h1_style.font.size.pt == 15
    assert h1_style.font.color.rgb == RGBColor(0, 0, 0)
    h1_rpr = h1_style.element.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr"
    )
    h1_fonts = h1_rpr.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts"
    )
    h1_color = h1_rpr.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color"
    )
    assert h1_fonts.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia") == "黑体"
    assert "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}asciiTheme" not in h1_fonts.attrib
    assert "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeColor" not in h1_color.attrib
    h2 = next(
        paragraph for paragraph in document.paragraphs if paragraph.text == "1.1 项目背景与意义"
    )
    assert h2.runs[0].font.size.pt == 14


def test_markdown_export_inserts_abstracts_after_title(db_session):
    project = Project(type="thesis", title="Hive 电商数仓", language="zh")
    chapter = Chapter(project=project, title="绪论", level=1, order=1)
    draft = ChapterDraft(
        chapter=chapter, version=1, content="# 绪论\n\n引言。", generation_mode="generate"
    )
    abstract = PaperAbstract(
        project=project,
        title_en="E-commerce Data Warehouse Based on Hive",
        abstract_zh="本文设计了电商数据仓库。",
        abstract_en="This paper designs a data warehouse.",
        keywords_zh=["Hive", "数据仓库"],
        keywords_en=["Hive", "data warehouse"],
    )
    db_session.add_all([project, chapter, draft, abstract])
    db_session.commit()

    markdown = build_markdown(
        project,
        [chapter],
        {chapter.id: draft},
        abstract,
    )

    assert markdown.index("# Hive 电商数仓") < markdown.index("# 摘要")
    assert markdown.index("# 摘要") < markdown.index("# Abstract")
    assert markdown.index("# Abstract") < markdown.index("# 绪论")
    assert "本文设计了电商数据仓库。" in markdown
    assert "关键词：Hive；数据仓库" in markdown
    assert "Keywords: Hive; data warehouse" in markdown
    assert "E-commerce Data Warehouse Based on Hive" in markdown


def test_docx_export_includes_zh_en_abstracts(tmp_path):
    project = Project(type="thesis", title="Hive 电商数仓", language="zh")
    chapter = Chapter(project=project, title="绪论", level=1, order=1)
    draft = ChapterDraft(
        chapter=chapter, version=1, content="# 绪论\n\n引言。", generation_mode="generate"
    )
    abstract = PaperAbstract(
        project=project,
        title_en="E-commerce Data Warehouse Based on Hive",
        abstract_zh="本文设计了电商数据仓库。",
        abstract_en="This paper designs a data warehouse.",
        keywords_zh=["Hive", "数据仓库"],
        keywords_en=["Hive", "data warehouse"],
    )
    output = tmp_path / "abstract.docx"

    build_docx(output, project, [chapter], {chapter.id: draft}, paper_abstract=abstract)

    texts = [paragraph.text for paragraph in Document(output).paragraphs]
    joined = "\n".join(texts)
    assert texts[0] == "Hive 电商数仓"
    assert "摘要" in texts
    assert "Abstract" in texts
    assert "本文设计了电商数据仓库。" in texts
    assert "关键词： Hive；数据仓库" in joined or "关键词：Hive；数据仓库" in joined
    assert "Keywords: Hive; data warehouse" in joined
    assert joined.index("摘要") < joined.index("Abstract") < joined.index("绪论")


def test_relation_listing_uses_chapter_hierarchy_order():
    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(engine)
    with Session(engine) as db_session:
        project = Project(type="thesis", title="Nested relations", language="zh")
        root = Chapter(project=project, title="Root", level=1, order=1)
        child = Chapter(project=project, parent=root, title="Child", level=2, order=1)
        later_root = Chapter(project=project, title="Later root", level=1, order=2)
        relations = [
            ChapterRelation(chapter=later_root, next_bridge="third"),
            ChapterRelation(chapter=child, next_bridge="second"),
            ChapterRelation(chapter=root, next_bridge="first"),
        ]
        db_session.add_all([project, root, child, later_root, *relations])
        db_session.commit()
        project_id = project.id

    def override_get_db():
        with Session(engine) as session:
            yield session

    app = create_app()
    app.dependency_overrides[get_db] = override_get_db
    client = TestClient(app)

    response = client.get(f"/api/projects/{project_id}/relations")

    assert response.status_code == 200
    assert [item["next_bridge"] for item in response.json()] == [
        "first",
        "second",
        "third",
    ]
