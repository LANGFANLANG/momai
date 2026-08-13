from docx import Document
from docx.oxml.ns import qn

from app.db.models import Chapter, ChapterDraft, Project
from app.export.academic import formula_label, to_chinese_num
from app.export.docx import build_docx
from app.export.docx_style import DocxStyle


def _build(tmp_path, content, chapters=None):
    project = Project(type="thesis", title="Hive 电商数仓", language="zh")
    if chapters is None:
        chapter = Chapter(project=project, title="系统设计", level=1, order=1)
        draft = ChapterDraft(
            chapter=chapter, version=1, content=content, generation_mode="generate"
        )
        chapters = [(chapter, draft)]
    drafts = {}
    chapter_objs = []
    for chapter, draft in chapters:
        chapter_objs.append(chapter)
        if draft is not None:
            drafts[chapter.id] = draft
    output = tmp_path / "academic.docx"
    build_docx(output, project, chapter_objs, drafts, DocxStyle())
    return Document(output)


def _border_sz(cell, edge: str) -> str | None:
    borders = cell._tc.find(qn("w:tcPr")).find(qn("w:tcBorders"))
    node = borders.find(qn(f"w:{edge}"))
    return node.get(qn("w:val")), node.get(qn("w:sz"))


def test_to_chinese_num_and_formula_label():
    assert to_chinese_num(1) == "一"
    assert to_chinese_num(10) == "十"
    assert to_chinese_num(12) == "十二"
    assert formula_label(1) == "（一）"


def test_docx_three_line_table_has_caption_and_borders(tmp_path):
    document = _build(
        tmp_path,
        (
            "本节给出字段说明。\n\n"
            "表 用户基本信息\n"
            "| 字段 | 类型 | 说明 |\n"
            "| --- | --- | --- |\n"
            "| user_id | INT | 用户标识 |\n"
            "| age | INT | 年龄 |\n"
        ),
    )
    texts = [paragraph.text for paragraph in document.paragraphs]
    assert "表1.1 用户基本信息" in texts
    assert document.tables
    table = document.tables[0]
    assert table.cell(0, 0).text.strip() == "字段"
    assert table.cell(1, 0).text.strip() == "user_id"
    top_val, top_sz = _border_sz(table.cell(0, 0), "top")
    header_bottom_val, header_bottom_sz = _border_sz(table.cell(0, 0), "bottom")
    last_bottom_val, last_bottom_sz = _border_sz(table.cell(2, 0), "bottom")
    left_val, _ = _border_sz(table.cell(1, 0), "left")
    mid_top_val, _ = _border_sz(table.cell(1, 0), "top")
    assert top_val == "single" and top_sz == "12"
    assert header_bottom_val == "single" and header_bottom_sz == "6"
    assert last_bottom_val == "single" and last_bottom_sz == "12"
    assert left_val == "nil"
    assert mid_top_val == "nil"
    joined = "\n".join(texts)
    assert "user_id    INT" not in joined


def test_docx_table_numbering_follows_chapter(tmp_path):
    project = Project(type="thesis", title="Hive 电商数仓", language="zh")
    first = Chapter(project=project, title="绪论", level=1, order=1)
    first.id = "chapter-1"
    first_draft = ChapterDraft(
        chapter=first,
        version=1,
        content="表 背景对比\n| 项目 | 说明 |\n| --- | --- |\n| A | 一 |\n",
        generation_mode="generate",
    )
    second = Chapter(project=project, title="系统设计", level=1, order=2)
    second.id = "chapter-2"
    second_draft = ChapterDraft(
        chapter=second,
        version=1,
        content=(
            "表 ODS层\n| 表名 | 说明 |\n| --- | --- |\n| ods_user | 用户 |\n\n"
            "表 DWD层\n| 表名 | 说明 |\n| --- | --- |\n| dwd_user | 明细 |\n"
        ),
        generation_mode="generate",
    )
    document = _build(
        tmp_path,
        "",
        chapters=[(first, first_draft), (second, second_draft)],
    )
    texts = [paragraph.text for paragraph in document.paragraphs]
    assert "表1.1 背景对比" in texts
    assert "表2.1 ODS层" in texts
    assert "表2.2 DWD层" in texts


def test_docx_display_formula_numbered_in_chinese(tmp_path):
    document = _build(
        tmp_path,
        (
            "准确率定义为：\n\n"
            "$$\n"
            "Accuracy = \\frac{TP + TN}{TP + TN + FP + FN}\n"
            "$$\n\n"
            "召回率为 $R = \\frac{TP}{TP+FN}$。\n\n"
            "$$ F_1 = 2 \\cdot \\frac{P R}{P + R} $$\n"
        ),
    )
    inline = [p for p in document.paragraphs if "召回率为" in p.text]
    assert inline
    omaths = document.element.body.findall(".//" + qn("m:oMath"))
    assert len(omaths) >= 3
    numbered = [
        table.cell(0, 2).text.strip()
        for table in document.tables
        if len(table.columns) == 3 and table.cell(0, 2).text.strip().startswith("（")
    ]
    assert numbered == ["（一）", "（二）"]
