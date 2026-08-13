import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from app.db.models import Chapter, ChapterDraft, PaperAbstract, Project, ProjectReference
from app.export.citations import format_reference_entry, is_bibliography_chapter, numbered_chapter_bodies
from app.export.common import (
    abstract_has_content,
    chapters_by_id,
    covered_by_ancestor_draft,
    join_keywords,
)
from app.export.academic import (
    AcademicCounters,
    caption_title,
    collect_table_rows,
    is_table_caption,
    is_table_row,
    next_formula_label,
    next_nonempty,
    next_table_caption,
    strip_formula_tag,
)
from app.export.formula import append_omath
from app.export.docx_style import DocxStyle

HEADING_RE = re.compile(r"^(#{1,6})\s*(.+?)\s*#*\s*$")
NUMBERED_HEADING_RE = re.compile(r"^\s*(\d+(?:\.\d+){0,5})(?:\s+|[.．、])")
UL_RE = re.compile(r"^[-*+]\s+(.+)$")
OL_RE = re.compile(r"^\d+[.)、]\s+(.+)$")
HR_RE = re.compile(r"^(-{3,}|\*{3,}|_{3,})$")
INLINE_RE = re.compile(
    r"(?<!\$)\$([^$\n]+)\$(?!\$)|"
    r"!\[[^\]]*\]\([^)]*\)|"
    r"\[([^\]]+)\]\([^)]*\)|"
    r"\*\*(.+?)\*\*|"
    r"__(.+?)__|"
    r"`([^`]+)`|"
    r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)|"
    r"~~(.+?)~~"
)
THEME_FONT_ATTRS = (
    qn("w:asciiTheme"),
    qn("w:hAnsiTheme"),
    qn("w:eastAsiaTheme"),
    qn("w:cstheme"),
)
THEME_COLOR_ATTRS = (qn("w:themeColor"), qn("w:themeTint"), qn("w:themeShade"))


def _apply_rpr(r_pr, ascii_font: str, east_asia: str, size_pt: float, bold: bool) -> None:
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:cs"), ascii_font)
    for attr in THEME_FONT_ATTRS:
        if attr in r_fonts.attrib:
            del r_fonts.attrib[attr]

    half_points = str(int(round(size_pt * 2)))
    for tag in ("w:sz", "w:szCs"):
        element = r_pr.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            r_pr.append(element)
        element.set(qn("w:val"), half_points)

    for tag, enabled in (("w:b", bold), ("w:bCs", bold)):
        element = r_pr.find(qn(tag))
        if enabled:
            if element is None:
                element = OxmlElement(tag)
                r_pr.append(element)
            element.set(qn("w:val"), "true")
        elif element is not None:
            r_pr.remove(element)

    for tag in ("w:i", "w:iCs"):
        element = r_pr.find(qn(tag))
        if element is not None:
            r_pr.remove(element)

    color = r_pr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        r_pr.append(color)
    color.set(qn("w:val"), "000000")
    for attr in THEME_COLOR_ATTRS:
        if attr in color.attrib:
            del color.attrib[attr]


def _set_run_font(run, ascii_font: str, east_asia: str, size_pt: float, bold: bool = False) -> None:
    _apply_rpr(run._element.get_or_add_rPr(), ascii_font, east_asia, size_pt, bold)


def _set_style_font(style_obj, ascii_font: str, east_asia: str, size_pt: float, bold: bool) -> None:
    _apply_rpr(style_obj.element.get_or_add_rPr(), ascii_font, east_asia, size_pt, bold)


def _apply_paragraph_spacing(paragraph, style: DocxStyle, *, indent_pt: float | None = None) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = style.line_spacing_multiple
    fmt.space_before = Pt(style.space_before_pt)
    fmt.space_after = Pt(style.space_after_pt)
    fmt.first_line_indent = Pt(0) if indent_pt is None else Pt(indent_pt)


def _configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


def _configure_styles(document: Document, style: DocxStyle) -> None:
    normal = document.styles["Normal"]
    _set_style_font(normal, style.body_ascii, style.body_east_asia, style.body_size_pt, False)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = style.line_spacing_multiple
    normal.paragraph_format.space_before = Pt(style.space_before_pt)
    normal.paragraph_format.space_after = Pt(style.space_after_pt)
    normal.paragraph_format.first_line_indent = Pt(style.body_size_pt * style.first_line_indent_chars)
    for level in range(1, 7):
        heading = document.styles[f"Heading {level}"]
        _set_style_font(
            heading,
            style.heading_ascii,
            style.heading_east_asia,
            _heading_size(style, level),
            True,
        )
        heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        heading.paragraph_format.line_spacing = style.line_spacing_multiple
        heading.paragraph_format.space_before = Pt(style.space_before_pt)
        heading.paragraph_format.space_after = Pt(style.space_after_pt)
        heading.paragraph_format.first_line_indent = Pt(0)


def _heading_size(style: DocxStyle, level: int) -> float:
    if level <= 1:
        return style.heading1_size_pt
    if level == 2:
        return style.heading2_size_pt
    return style.heading3_size_pt


def _word_heading_level(markdown_level: int, text: str) -> int:
    numbered = NUMBERED_HEADING_RE.match(_plain_text(text))
    if numbered:
        return min(numbered.group(1).count(".") + 1, 6)
    return markdown_level


def _clean_residual(text: str) -> str:
    text = re.sub(r"<!--.*?-->", "", text)
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"!\[([^\]]*)\]\([^)]*\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)
    return text.replace("**", "").replace("__", "").replace("`", "")


def _inline_chunks(text: str):
    text = re.sub(r"<!--.*?-->", "", text)
    text = re.sub(r"<[^>]+>", "", text)
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            yield _clean_residual(text[pos : match.start()]), "plain"
        if match.group(0).startswith("$"):
            yield match.group(1), "math"
        elif match.group(0).startswith("!["):
            alt = re.match(r"!\[([^\]]*)\]", match.group(0))
            yield (alt.group(1) if alt else ""), "plain"
        elif match.group(2):
            yield match.group(2), "plain"
        elif match.group(3) or match.group(4):
            yield match.group(3) or match.group(4), "bold"
        else:
            yield match.group(5) or match.group(6) or match.group(7) or "", "plain"
        pos = match.end()
    if pos < len(text) or pos == 0:
        yield _clean_residual(text[pos:]), "plain"


def _plain_text(text: str) -> str:
    return "".join(chunk for chunk, _ in _inline_chunks(text))


def _add_inline_runs(paragraph, text: str, ascii_font: str, east_asia: str, size_pt: float, *, base_bold: bool) -> None:
    emitted = False
    for chunk, kind in _inline_chunks(text):
        if kind == "math":
            append_omath(paragraph, chunk)
            emitted = True
            continue
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        _set_run_font(run, ascii_font, east_asia, size_pt, base_bold or kind == "bold")
        emitted = True
    if not emitted:
        run = paragraph.add_run("")
        _set_run_font(run, ascii_font, east_asia, size_pt, base_bold)


def _add_title(document: Document, text: str, style: DocxStyle) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_paragraph_spacing(paragraph, style)
    _add_inline_runs(
        paragraph,
        text,
        style.heading_ascii,
        style.heading_east_asia,
        style.heading1_size_pt,
        base_bold=True,
    )


def _add_heading(document: Document, text: str, level: int, style: DocxStyle) -> None:
    style_level = min(max(level, 1), 6)
    paragraph = document.add_paragraph()
    paragraph.style = document.styles[f"Heading {style_level}"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _apply_paragraph_spacing(paragraph, style)
    _add_inline_runs(
        paragraph,
        _plain_text(text),
        style.heading_ascii,
        style.heading_east_asia,
        _heading_size(style, level),
        base_bold=True,
    )


def _add_body(document: Document, text: str, style: DocxStyle, *, indent: bool = True) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    indent_pt = style.body_size_pt * style.first_line_indent_chars if indent else 0
    _apply_paragraph_spacing(paragraph, style, indent_pt=indent_pt)
    _add_inline_runs(
        paragraph,
        text,
        style.body_ascii,
        style.body_east_asia,
        style.body_size_pt,
        base_bold=False,
    )


def _nil_border(edge: str, sz: int = 0):
    element = OxmlElement(f"w:{edge}")
    if sz:
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(sz))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
    else:
        element.set(qn("w:val"), "nil")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "auto")
    return element


def _set_cell_borders(cell, top: int = 0, bottom: int = 0) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcBorders"))
    if existing is not None:
        tc_pr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    borders.append(_nil_border("top", top))
    borders.append(_nil_border("left", 0))
    borders.append(_nil_border("bottom", bottom))
    borders.append(_nil_border("right", 0))
    tc_pr.append(borders)


def _set_tbl_width_pct(table, pct: int = 5000) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(pct))
    tbl_w.set(qn("w:type"), "pct")
    existing_jc = tbl_pr.find(qn("w:jc"))
    if existing_jc is not None:
        tbl_pr.remove(existing_jc)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    tbl_pr.append(jc)
    existing_borders = tbl_pr.find(qn("w:tblBorders"))
    if existing_borders is not None:
        tbl_pr.remove(existing_borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(_nil_border(edge, 0))
    tbl_pr.append(borders)


def _set_cell_width_pct(cell, pct: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(pct))
    tc_w.set(qn("w:type"), "pct")


def _fill_cell(cell, text: str, style: DocxStyle, *, bold: bool, size_pt: float) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_paragraph_spacing(paragraph, style, indent_pt=0)
    _add_inline_runs(
        paragraph,
        text,
        style.body_ascii,
        style.body_east_asia,
        size_pt,
        base_bold=bold,
    )


def _add_table_caption(document: Document, caption: str, style: DocxStyle) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_paragraph_spacing(paragraph, style, indent_pt=0)
    _add_inline_runs(
        paragraph,
        caption,
        style.heading_ascii,
        style.heading_east_asia,
        10.5,
        base_bold=False,
    )


def _add_three_line_table(
    document: Document,
    rows: list[list[str]],
    title: str,
    style: DocxStyle,
    counters: AcademicCounters,
) -> None:
    if not rows:
        return
    _add_table_caption(document, next_table_caption(counters, title), style)
    cols = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=cols)
    table.autofit = True
    _set_tbl_width_pct(table)
    size_pt = min(style.body_size_pt, 10.5)
    last = len(rows) - 1
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            value = row[c_idx] if c_idx < len(row) else ""
            top = 12 if r_idx == 0 else 0
            bottom = 0
            if r_idx == 0:
                bottom = 6 if last > 0 else 12
            if r_idx == last:
                bottom = 12
            _set_cell_borders(cell, top=top, bottom=bottom)
            _fill_cell(cell, _plain_text(value), style, bold=r_idx == 0, size_pt=size_pt)


def _add_display_formula(
    document: Document,
    latex: str,
    style: DocxStyle,
    counters: AcademicCounters,
) -> None:
    latex = strip_formula_tag(latex)
    if not latex:
        return
    label = next_formula_label(counters)
    table = document.add_table(rows=1, cols=3)
    table.autofit = True
    _set_tbl_width_pct(table)
    widths = (750, 3500, 750)
    alignments = (
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.RIGHT,
    )
    for index, cell in enumerate(table.rows[0].cells):
        _set_cell_borders(cell)
        _set_cell_width_pct(cell, widths[index])
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = alignments[index]
        _apply_paragraph_spacing(paragraph, style, indent_pt=0)
        if index == 1:
            append_omath(paragraph, latex)
        elif index == 2:
            run = paragraph.add_run(label)
            _set_run_font(run, style.body_ascii, style.body_east_asia, style.body_size_pt)


def _collect_display_formula(lines: list[str], start: int) -> tuple[int, str]:
    stripped = lines[start].strip()
    if stripped.startswith("$$") and stripped.endswith("$$") and len(stripped) > 4:
        return start + 1, stripped[2:-2].strip()
    if stripped.startswith("\\[") and stripped.endswith("\\]") and len(stripped) > 4:
        return start + 1, stripped[2:-2].strip()
    opener = "$$" if stripped.startswith("$$") else "\\["
    closer = "$$" if opener == "$$" else "\\]"
    parts = [stripped[len(opener):]] if stripped != opener else []
    index = start + 1
    while index < len(lines):
        current = lines[index].strip()
        if current.endswith(closer):
            parts.append(current[: -len(closer)])
            return index + 1, "\n".join(parts).strip()
        parts.append(lines[index])
        index += 1
    return index, "\n".join(parts).strip()


def _add_markdown(
    document: Document,
    markdown: str,
    style: DocxStyle,
    counters: AcademicCounters | None = None,
) -> None:
    counters = counters or AcademicCounters(chapter_no=1)
    lines = markdown.splitlines()
    in_code = False
    index = 0
    while index < len(lines):
        stripped = lines[index].strip()
        if stripped.startswith("```"):
            in_code = not in_code
            index += 1
            continue
        if not stripped or HR_RE.fullmatch(stripped):
            index += 1
            continue
        if in_code:
            _add_body(document, stripped, style, indent=False)
            index += 1
            continue
        if stripped.startswith(">"):
            stripped = stripped.lstrip("> ").strip()
            if not stripped:
                index += 1
                continue
        if stripped.startswith("$$") or stripped.startswith("\\["):
            index, latex = _collect_display_formula(lines, index)
            _add_display_formula(document, latex, style, counters)
            continue
        if is_table_caption(stripped) and is_table_row(next_nonempty(lines, index + 1) or ""):
            title = caption_title(stripped)
            index, rows = collect_table_rows(lines, index + 1)
            _add_three_line_table(document, rows, title, style, counters)
            continue
        if is_table_row(stripped):
            index, rows = collect_table_rows(lines, index)
            _add_three_line_table(document, rows, "", style, counters)
            continue
        heading = HEADING_RE.match(stripped)
        if heading:
            text = heading.group(2).strip()
            level = _word_heading_level(len(heading.group(1)), text)
            _add_heading(document, text, level, style)
            index += 1
            continue
        bullet = UL_RE.match(stripped)
        if bullet:
            _add_body(document, f"• {bullet.group(1).strip()}", style, indent=False)
            index += 1
            continue
        numbered = OL_RE.match(stripped)
        if numbered:
            _add_body(document, stripped, style, indent=False)
            index += 1
            continue
        _add_body(document, stripped, style)
        index += 1


def _add_centered(document: Document, text: str, style: DocxStyle, size_pt: float) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_paragraph_spacing(paragraph, style)
    _add_inline_runs(
        paragraph,
        text,
        style.heading_ascii,
        style.heading_east_asia,
        size_pt,
        base_bold=True,
    )


def _add_abstracts(document: Document, abstract: PaperAbstract | None, style: DocxStyle) -> None:
    if not abstract_has_content(abstract) or abstract is None:
        return
    if (abstract.abstract_zh or "").strip():
        _add_heading(document, "摘要", 1, style)
        _add_markdown(document, abstract.abstract_zh, style)
        keywords = join_keywords(abstract.keywords_zh, chinese=True)
        if keywords:
            _add_body(document, f"**关键词：** {keywords}", style, indent=False)
    if (abstract.abstract_en or "").strip() or (abstract.title_en or "").strip():
        _add_heading(document, "Abstract", 1, style)
        if (abstract.title_en or "").strip():
            _add_centered(document, abstract.title_en.strip(), style, style.heading2_size_pt)
        if (abstract.abstract_en or "").strip():
            _add_markdown(document, abstract.abstract_en, style)
        keywords = join_keywords(abstract.keywords_en, chinese=False)
        if keywords:
            _add_body(document, f"**Keywords:** {keywords}", style, indent=False)


def build_docx(
    output_path: Path,
    project: Project,
    chapters: list[Chapter],
    drafts: dict[str, ChapterDraft],
    style: DocxStyle | None = None,
    paper_abstract: PaperAbstract | None = None,
    references: list[ProjectReference] | None = None,
) -> None:
    style = style or DocxStyle()
    document = Document()
    _configure_page(document)
    _configure_styles(document, style)
    _add_title(document, project.title, style)
    _add_abstracts(document, paper_abstract, style)
    by_id = chapters_by_id(chapters)
    bodies, bibliography = numbered_chapter_bodies(chapters, drafts, references)
    counters = AcademicCounters()
    for chapter in chapters:
        if is_bibliography_chapter(chapter.title) or covered_by_ancestor_draft(chapter, by_id, drafts):
            continue
        if chapter.level <= 1:
            counters.chapter_no += 1
            counters.table_no = 0
        _add_heading(document, chapter.title, max(chapter.level, 1), style)
        content = bodies.get(chapter.id)
        if content:
            _add_markdown(document, content, style, counters)
    if bibliography:
        _add_heading(document, "参考文献", 1, style)
        for number, ref in bibliography:
            _add_body(document, format_reference_entry(ref, number), style, indent=False)
    document.save(output_path)
